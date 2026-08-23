#!/usr/bin/env python3
"""Build the AI Week representative-case submission report DOCX."""

from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "보고서 산출물/AI_Week_Promo_Web_Builder_대표사례_제출보고서.md"
OUTPUT = ROOT / "보고서 산출물/AI_Week_Promo_Web_Builder_대표사례_제출보고서.docx"
BASE_SCRIPT = ROOT / "scripts/build_agenda_progress_report_docx.py"

spec = importlib.util.spec_from_file_location("agenda_report", BASE_SCRIPT)
base = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(base)
ORIGINAL_TABLE_WEIGHTS = base.table_weights
ORIGINAL_ADD_TABLE = base.add_table


def replace_header_label(doc):
    for table in doc.sections[0].header.tables:
        for cell in table._element.xpath(".//w:tc"):
            for text in cell.xpath(".//w:t"):
                if text.text == "종합 진행·마일스톤 보고서":
                    text.text = "AI WEEK · 대표 사례 제출 보고서"


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    base.set_font(p.add_run("AI WEEK 대표 사례"), size=23, bold=True, color=base.INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    base.set_font(
        p.add_run("Promo Web Builder | 프로모션 제작 핸드오프 End-to-End 자동화"),
        size=14,
        color=base.MUTED,
    )

    metadata = [
        ("핵심 병목", "요구사항이 문서·이미지·레이아웃·코드로 반복 변환되는 조직 간 핸드오프"),
        ("연결 구조", "자연어→AI Builder→Composition→Asset→Visual Editor→Quality Gate→Web Output"),
        ("확인 수치", "자동 테스트 0→137개, Desktop·Mobile 품질 검사 및 Browser E2E 도입"),
        ("실현 단계", "구축 중 — 핵심 구현·로컬 검증 완료, 최신분 Production 배포 전"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        base.set_font(p.add_run(f"{label}: "), size=10.2, bold=True, color=base.INK)
        base.set_font(p.add_run(value), size=10.2, color=RGBColor(0x22, 0x22, 0x22))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(13)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    callout = doc.add_table(rows=1, cols=1)
    base.apply_table_geometry(callout, [base.CONTENT_WIDTH_DXA], indent_dxa=120)
    cell = callout.cell(0, 0)
    base.shade_cell(cell, base.ACCENT_FILL)
    base.set_cell_border(cell, color="C9D7E6", size="4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    base.set_font(p.add_run("핵심 답변  "), size=10.5, bold=True, color=base.DARK_BLUE)
    base.set_font(
        p.add_run(
            "기획 산출물이 다음 조직의 편집·검증·배포 인풋으로 자동 변환되도록 "
            "AI 생성, 자산 준비, 품질 게이트를 하나의 문서 계약으로 연결했다."
        ),
        size=10.5,
        color=base.INK,
    )


def add_properties(doc):
    props = doc.core_properties
    props.title = "AI Week 대표 사례 — Promo Web Builder"
    props.subject = "프로모션 제작 조직 간 핸드오프 End-to-End 자동화"
    props.author = "Promo Web Builder Project"
    props.keywords = "AI Week, AI Agent, Automation, Promo Web Builder, Quality Gate"


def apply_visual_section_breaks(doc):
    """Keep major infographic-led sections intact and prevent top-edge heading collisions."""
    break_prefixes = (
        "2. 개선 프로세스",
        "4. 연결 구조",
        "5. 정량적 효율화 수준",
        "7. 실현 단계",
    )
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(break_prefixes):
            paragraph.paragraph_format.page_break_before = True


def submission_table_weights(rows):
    """Named table-width overrides for long evidence and completion-condition columns."""
    cols = len(rows[0])
    if rows[0] == ["구축 완료", "로컬 검증 완료", "운영 측정 필요", "확장 후보"]:
        return [1.625, 1.625, 1.625, 1.625]
    if rows[0] == ["흐름", "노드", "수행 역할", "다음 노드로 전달되는 인풋"]:
        return [0.55, 1.55, 2.70, 1.70]
    if cols == 4:
        return [0.95, 1.60, 2.65, 1.30]
    if cols == 5:
        return [1.10, 0.82, 1.92, 1.08, 1.58]
    if cols == 3 and rows[0] in (
        ["As-Is 병목", "AI·자동화 연결", "To-Be 결과"],
        ["자동 테스트", "품질 검사 범위", "중간 핸드오프"],
    ):
        return [2.20, 1.85, 2.45]
    return ORIGINAL_TABLE_WEIGHTS(rows)


def style_cell_text(cell, *, size=None, bold=None, color=None, align=None):
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            base.set_font(run, size=size, bold=bold, color=color)


def submission_add_table(doc, rows):
    """Render selected comparison/status tables as native Word infographics."""
    ORIGINAL_ADD_TABLE(doc, rows)
    table = doc.tables[-1]
    header = rows[0]

    if header == ["As-Is 병목", "AI·자동화 연결", "To-Be 결과"]:
        fills = ["F2F4F7", "E8EEF5", "EAF4EE"]
        for col, cell in enumerate(table.rows[0].cells):
            base.shade_cell(cell, "1F4D78")
            style_cell_text(cell, size=9.2, bold=True, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for row in table.rows[1:]:
            for col, cell in enumerate(row.cells):
                base.shade_cell(cell, fills[col])
                style_cell_text(cell, size=8.8, align=WD_ALIGN_PARAGRAPH.LEFT)

    elif header == ["자동 테스트", "품질 검사 범위", "중간 핸드오프"]:
        for cell in table.rows[0].cells:
            base.shade_cell(cell, "1F4D78")
            style_cell_text(cell, size=9.5, bold=True, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in table.rows[1].cells:
            base.shade_cell(cell, "E8EEF5")
            style_cell_text(cell, size=16, bold=True, color=base.DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in table.rows[2].cells:
            base.shade_cell(cell, "F4F6F9")
            style_cell_text(cell, size=8.6, color=base.INK, align=WD_ALIGN_PARAGRAPH.CENTER)

    elif header == ["구축 완료", "로컬 검증 완료", "운영 측정 필요", "확장 후보"]:
        header_fills = ["1F4D78", "2E74B5", "7A5A00", "666B73"]
        body_fills = ["E8EEF5", "EAF1F8", "FFF5D6", "F2F4F7"]
        for col, cell in enumerate(table.rows[0].cells):
            base.shade_cell(cell, header_fills[col])
            style_cell_text(cell, size=8.8, bold=True, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for row in table.rows[1:]:
            for col, cell in enumerate(row.cells):
                base.shade_cell(cell, body_fills[col])
                style_cell_text(cell, size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER)

    elif header == ["흐름", "노드", "수행 역할", "다음 노드로 전달되는 인풋"]:
        for cell in table.rows[0].cells:
            base.shade_cell(cell, "1F4D78")
            style_cell_text(cell, size=9, bold=True, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for row_index, row in enumerate(table.rows[1:], 1):
            base.shade_cell(row.cells[0], "2E74B5")
            style_cell_text(row.cells[0], size=10, bold=True, color=base.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            for cell in row.cells[1:]:
                base.shade_cell(cell, "F4F6F9" if row_index % 2 else "FFFFFF")


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    base.configure_document(doc)
    replace_header_label(doc)
    add_properties(doc)
    add_masthead(doc)
    base.table_weights = submission_table_weights
    base.add_table = submission_add_table
    base.add_markdown_content(doc, markdown)
    apply_visual_section_breaks(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
