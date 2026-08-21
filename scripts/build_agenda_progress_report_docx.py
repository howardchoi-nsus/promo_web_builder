#!/usr/bin/env python3
"""Build the agenda-based progress report DOCX from the reviewed Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/자료/promo-web-builder-agenda-progress-report-2026-06-08-to-2026-08-21.md"
OUTPUT = ROOT / "docs/자료/Promo_Web_Builder_Agenda_Progress_Report_2026-06-08_to_2026-08-21.docx"
SKILL_DIR = Path(
    "/Users/hojunchoi/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.819.11345/skills/documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x66, 0x6B, 0x73)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
ACCENT_FILL = "E8EEF5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# Named override for Korean glyph coverage in Word and LibreOffice renderers.
FONT = "Noto Sans CJK KR"
EAST_ASIA_FONT = "Noto Sans CJK KR"
CONTENT_WIDTH_DXA = 9360


def set_font(run, *, size=None, bold=None, color=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DBE2", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_repeat_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field_run = paragraph.add_run()
    set_font(field_run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    field_run._r.append(fld_begin)
    instr_run = paragraph.add_run()
    set_font(instr_run, size=9, color=MUTED)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    sep_run = paragraph.add_run()
    set_font(sep_run, size=9, color=MUTED)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)
    value_run = paragraph.add_run("1")
    set_font(value_run, size=9, color=MUTED)
    end_run = paragraph.add_run()
    set_font(end_run, size=9, color=MUTED)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def new_decimal_num_id(doc):
    """Create a fresh real decimal numbering instance so each process restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for level in abstract.findall(qn("w:lvl")):
            if level.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = level.find(qn("w:numFmt"))
            if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
                abstract_id = abstract.get(qn("w:abstractNumId"))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        raise RuntimeError("No decimal numbering definition found")
    existing = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_num_id(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_ref])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    widths = column_widths_from_weights([3.8, 2.7], CONTENT_WIDTH_DXA)
    apply_table_geometry(table, widths, indent_dxa=120)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(left.add_run("PROMO WEB BUILDER"), size=8.5, bold=True, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(right.add_run("종합 진행·마일스톤 보고서"), size=8.5, color=MUTED)
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "D7DBE2")
        borders.append(bottom)
        tc_pr.append(borders)

    footer = section.footer
    p = footer.paragraphs[0]
    set_repeat_page_number(p)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("개발 종합 진행 보고서"), size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("Promo Web Builder | 종합 진행·마일스톤·Agenda 기준"), size=14, color=MUTED)

    metadata = [
        ("보고 기간", "2026년 6월 8일 ~ 2026년 8월 21일"),
        ("작성 기준일", "2026년 8월 21일"),
        ("보고 대상", "Promo Web Builder 개발 및 품질 고도화"),
        ("보고 목적", "종합 진행 현황, 마일스톤, Agenda별 목표·과정·결과 산출물 보고"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=INK)
        set_font(p.add_run(value), size=10.5, color=RGBColor(0x22, 0x22, 0x22))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    callout = doc.add_table(rows=1, cols=1)
    apply_table_geometry(callout, [CONTENT_WIDTH_DXA], indent_dxa=120)
    cell = callout.cell(0, 0)
    shade_cell(cell, ACCENT_FILL)
    set_cell_border(cell, color="C9D7E6", size="4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("핵심 요약  "), size=10.5, bold=True, color=DARK_BLUE)
    set_font(
        p.add_run(
            "초기 AI UI 이미지 생성 POC를 Registry 기반 Composition, 공통 Visual Editor, "
            "Asset Readiness, Desktop/Mobile Quality Gate를 갖춘 제작 플랫폼으로 전환했다."
        ),
        size=10.5,
        color=INK,
    )


def add_inline_runs(paragraph, text, *, base_size=11, color=None, bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9, color=RGBColor(0x44, 0x44, 0x44))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size, color=color, bold=bold)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, index


def table_weights(rows):
    cols = len(rows[0])
    lengths = []
    for col in range(cols):
        maximum = max(len(re.sub(r"[`*]", "", row[col])) for row in rows if col < len(row))
        lengths.append(max(5, min(maximum, 30)))
    if cols == 2:
        return [1.35, 4.65]
    if cols == 3:
        if rows[0] == ["구분", "초기 상태", "현재 결과"]:
            return [1.5, 3.2, 1.8]
        return [1.2, 3.7, 1.6]
    if cols == 4:
        return [1.2, 1.9, 2.6, 0.8]
    if cols == 5:
        return [1.25, 1.55, 2.15, 1.1, 0.55]
    return lengths


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = column_widths_from_weights(table_weights(rows), CONTENT_WIDTH_DXA)
    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        if row_index == 0:
            repeat_table_header(row)
        for col_index, text in enumerate(row_data):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(cell, LIGHT_FILL if row_index == 0 else "FFFFFF")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            if col_index == cols - 1 and len(text) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(
                p,
                text,
                base_size=8.4 if cols >= 4 else 9,
                color=INK if row_index == 0 else RGBColor(0x22, 0x22, 0x22),
                bold=row_index == 0,
            )
    apply_table_geometry(table, widths, indent_dxa=120)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    add_inline_runs(p, text, base_size=10, color=INK)


def add_markdown_content(doc, markdown):
    lines = markdown.splitlines()
    index = 0
    title_skipped = False
    active_decimal_num_id = None
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line or line == "---":
            if line == "---":
                active_decimal_num_id = None
            index += 1
            continue
        if line.startswith("|"):
            active_decimal_num_id = None
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not title_skipped:
                title_skipped = True
                index += 1
                continue
            if text == "6. Agenda별 상세 진행 및 결과":
                index += 1
                continue
            if text.startswith("Agenda "):
                level = 1
            active_decimal_num_id = None
            style_level = min(max(level - 1, 1), 3)
            p = doc.add_paragraph(style=f"Heading {style_level}")
            add_inline_runs(p, text, base_size={1: 16, 2: 13, 3: 12}[style_level], bold=True)
            set_keep_with_next(p)
            index += 1
            continue
        if line.startswith(">"):
            active_decimal_num_id = None
            add_note(doc, line[1:].strip())
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            if active_decimal_num_id is None:
                active_decimal_num_id = new_decimal_num_id(doc)
            p = doc.add_paragraph(style="List Number")
            apply_num_id(p, active_decimal_num_id)
            add_inline_runs(p, numbered.group(1))
            index += 1
            continue
        if line.startswith("- "):
            active_decimal_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            index += 1
            continue
        p = doc.add_paragraph()
        active_decimal_num_id = None
        add_inline_runs(p, line)
        index += 1


def add_document_properties(doc):
    props = doc.core_properties
    props.title = "Promo Web Builder 개발 종합 진행 및 마일스톤 보고서"
    props.subject = "2026-06-08 ~ 2026-08-21 개발 진행 및 결과"
    props.author = "Promo Web Builder Project"
    props.keywords = "Promo Web Builder, AI Builder, Visual Editor, Quality Gate"


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_document_properties(doc)
    add_masthead(doc)
    add_markdown_content(doc, markdown)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
